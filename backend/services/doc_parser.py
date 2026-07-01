from pathlib import Path

import fitz
from docx import Document

from models.schemas import ParsedDocument

SUPPORTED_TYPES = {".pdf", ".docx", ".txt", ".md"}


def parse_document(file_path: str, filename: str) -> ParsedDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type: {suffix}")

    if suffix == ".pdf":
        return _parse_pdf(file_path, filename)
    elif suffix == ".docx":
        return _parse_docx(file_path, filename)
    else:
        return _parse_text(file_path, filename, suffix)


def _parse_pdf(file_path: str, filename: str) -> ParsedDocument:
    doc = fitz.open(file_path)

    if doc.is_encrypted:
        doc.close()
        raise ValueError("Cannot process encrypted PDF")

    pages = []
    all_text = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        pages.append({"page_number": page_num + 1, "text": text})
        all_text.append(text)

    full_text = "\n".join(all_text)
    total_pages = len(doc)
    doc.close()

    if not full_text.strip():
        raise ValueError("PDF contains no extractable text")

    return ParsedDocument(
        filename=filename,
        file_type="pdf",
        total_pages=total_pages,
        total_characters=len(full_text),
        text_content=full_text,
        pages=pages,
    )


def _parse_docx(file_path: str, filename: str) -> ParsedDocument:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    if not full_text.strip():
        raise ValueError("DOCX file contains no text")

    return ParsedDocument(
        filename=filename,
        file_type="docx",
        total_pages=None,
        total_characters=len(full_text),
        text_content=full_text,
        pages=None,
    )


def _parse_text(file_path: str, filename: str, suffix: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8") as f:
        full_text = f.read()

    if not full_text.strip():
        raise ValueError("File is empty")

    file_type = "md" if suffix == ".md" else "txt"

    return ParsedDocument(
        filename=filename,
        file_type=file_type,
        total_pages=None,
        total_characters=len(full_text),
        text_content=full_text,
        pages=None,
    )
